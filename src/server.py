import json
import io
import os
import base64
from typing import Optional, List, Dict, Any
from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from azure.storage.blob import BlobServiceClient, BlobClient
from PIL import Image

# Initialize FastMCP server
server = FastMCP("docx_mcp_formatting")

def get_blob_client(blob_uri: str) -> BlobClient:
    """Get Azure Blob client from URI"""
    return BlobClient.from_blob_url(blob_uri)

def download_docx_from_blob(blob_uri: str) -> Document:
    """Download DOCX file from Azure Blob Storage"""
    blob_client = get_blob_client(blob_uri)
    blob_data = blob_client.download_blob()
    docx_bytes = blob_data.readall()
    return Document(io.BytesIO(docx_bytes))

def upload_docx_to_blob(doc: Document, blob_uri: str) -> str:
    """Upload DOCX file to Azure Blob Storage"""
    blob_client = get_blob_client(blob_uri)
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    blob_client.upload_blob(docx_buffer, overwrite=True)
    return f"Document uploaded successfully to {blob_uri}"

@server.tool()
async def create_docx(
    blob_uri: str,
    title: Optional[str] = None,
    initial_content: Optional[str] = None
) -> str:
    """Create a new DOCX file and upload to Azure Blob Storage.

    Args:
        blob_uri: Azure Blob Storage URI where the document will be created
        title: Optional title to add as first heading
        initial_content: Optional initial text content
    
    Returns:
        Success message with document details
    """
    try:
        doc = Document()
        
        if title:
            doc.add_heading(title, 0)
        
        if initial_content:
            doc.add_paragraph(initial_content)
        
        result = upload_docx_to_blob(doc, blob_uri)
        return json.dumps({
            "success": True,
            "message": result,
            "blob_uri": blob_uri
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)

@server.tool()
async def read_docx(blob_uri: str) -> str:
    """Read and extract content from a DOCX file in Azure Blob Storage.

    Args:
        blob_uri: Azure Blob Storage URI of the document to read
    
    Returns:
        JSON with document structure including paragraphs, tables, and formatting
    """
    try:
        doc = download_docx_from_blob(blob_uri)
        
        content = {
            "paragraphs": [],
            "tables": [],
            "sections": []
        }
        
        # Extract paragraphs with formatting
        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name,
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": []
            }
            
            for run in para.runs:
                run_info = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "font_name": run.font.name
                }
                para_info["runs"].append(run_info)
            
            content["paragraphs"].append(para_info)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "table_index": table_idx,
                "rows": []
            }
            
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data["rows"].append(row_data)
            
            content["tables"].append(table_data)
        
        # Extract section information (page margins, etc.)
        for section in doc.sections:
            section_info = {
                "page_width": section.page_width.inches if section.page_width else None,
                "page_height": section.page_height.inches if section.page_height else None,
                "left_margin": section.left_margin.inches if section.left_margin else None,
                "right_margin": section.right_margin.inches if section.right_margin else None,
                "top_margin": section.top_margin.inches if section.top_margin else None,
                "bottom_margin": section.bottom_margin.inches if section.bottom_margin else None
            }
            content["sections"].append(section_info)
        
        return json.dumps({
            "success": True,
            "blob_uri": blob_uri,
            "content": content
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)

@server.tool()
async def insert_text(
    blob_uri: str,
    text: str,
    position: str = "end",
    paragraph_index: Optional[int] = None,
    style: Optional[str] = None
) -> str:
    """Insert text into a DOCX file at specified position.

    Args:
        blob_uri: Azure Blob Storage URI of the document
        text: Text content to insert
        position: Where to insert - "start", "end", or "at_index" (default: "end")
        paragraph_index: Paragraph index when position is "at_index"
        style: Optional paragraph style (e.g., "Heading 1", "Normal")
    
    Returns:
        Success message
    """
    try:
        doc = download_docx_from_blob(blob_uri)
        
        if position == "start":
            # Insert at beginning
            para = doc.paragraphs[0].insert_paragraph_before(text) if doc.paragraphs else doc.add_paragraph(text)
        elif position == "at_index" and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index].insert_paragraph_before(text)
            else:
                para = doc.add_paragraph(text)
        else:  # end
            para = doc.add_paragraph(text)
        
        if style:
            para.style = style
        
        result = upload_docx_to_blob(doc, blob_uri)
        return json.dumps({
            "success": True,
            "message": "Text inserted successfully",
            "blob_uri": blob_uri
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)

@server.tool()
async def edit_text(
    blob_uri: str,
    paragraph_index: int,
    new_text: str,
    preserve_formatting: bool = True
) -> str:
    """Edit existing text in a specific paragraph.

    Args:
        blob_uri: Azure Blob Storage URI of the document
        paragraph_index: Index of the paragraph to edit (0-based)
        new_text: New text content
        preserve_formatting: Whether to preserve existing formatting (default: True)
    
    Returns:
        Success message
    """
    try:
        doc = download_docx_from_blob(blob_uri)
        
        if paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                "success": False,
                "error": f"Paragraph index {paragraph_index} out of range. Document has {len(doc.paragraphs)} paragraphs."
            }, ensure_ascii=False)
        
        para = doc.paragraphs[paragraph_index]
        
        if preserve_formatting and para.runs:
            # Keep formatting of first run
            first_run = para.runs[0]
            bold = first_run.bold
            italic = first_run.italic
            font_name = first_run.font.name
            font_size = first_run.font.size
            
            # Clear paragraph and add new text with formatting
            para.clear()
            new_run = para.add_run(new_text)
            new_run.bold = bold
            new_run.italic = italic
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
        else:
            para.text = new_text
        
        result = upload_docx_to_blob(doc, blob_uri)
        return json.dumps({
            "success": True,
            "message": f"Paragraph {paragraph_index} edited successfully",
            "blob_uri": blob_uri
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)

@server.tool()
async def insert_or_edit_table(
    blob_uri: str,
    table_data: List[List[str]],
    table_index: Optional[int] = None,
    position: str = "end",
    style: Optional[str] = "Table Grid"
) -> str:
    """Insert a new table or edit an existing table in the document.

    Args:
        blob_uri: Azure Blob Storage URI of the document
        table_data: 2D array of table data [[row1_col1, row1_col2], [row2_col1, row2_col2]]
        table_index: Index of existing table to edit (None to insert new)
        position: Where to insert new table - "start", "end" (default: "end")
        style: Table style name (default: "Table Grid")
    
    Returns:
        Success message
    """
    try:
        doc = download_docx_from_blob(blob_uri)
        
        if not table_data or not table_data[0]:
            return json.dumps({
                "success": False,
                "error": "Table data cannot be empty"
            }, ensure_ascii=False)
        
        rows = len(table_data)
        cols = len(table_data[0])
        
        if table_index is not None and table_index < len(doc.tables):
            # Edit existing table
            table = doc.tables[table_index]
            
            # Add rows if needed
            while len(table.rows) < rows:
                table.add_row()
            
            # Update cells
            for row_idx, row_data in enumerate(table_data):
                if row_idx < len(table.rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_data
            
            message = f"Table {table_index} edited successfully"
        else:
            # Insert new table
            table = doc.add_table(rows=rows, cols=cols)
            
            # Set table style
            if style:
                table.style = style
            
            # Populate table
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_data in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = cell_data
            
            message = "New table inserted successfully"
        
        result = upload_docx_to_blob(doc, blob_uri)
        return json.dumps({
            "success": True,
            "message": message,
            "blob_uri": blob_uri
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)

@server.tool()
async def insert_or_edit_image(
    blob_uri: str,
    image_data: str,
    width_inches: float = 3.0,
    image_index: Optional[int] = None,
    position: str = "end"
) -> str:
    """Insert or replace an image in the document.

    Args:
        blob_uri: Azure Blob Storage URI of the document
        image_data: Base64 encoded image data or Azure Blob URI of image
        width_inches: Width of image in inches (default: 3.0)
        image_index: Index of existing image to replace (None to insert new)
        position: Where to insert - "start", "end" (default: "end")
    
    Returns:
        Success message
    """
    try:
        doc = download_docx_from_blob(blob_uri)
        
        # Handle image data
        if image_data.startswith('http'):
            # Download from Azure Blob
            image_blob_client = get_blob_client(image_data)
            image_bytes = image_blob_client.download_blob().readall()
            image_stream = io.BytesIO(image_bytes)
        else:
            # Decode base64
            image_bytes = base64.b64decode(image_data)
            image_stream = io.BytesIO(image_bytes)
        
        # Add image to document
        if position == "start" and doc.paragraphs:
            para = doc.paragraphs[0].insert_paragraph_before()
        else:
            para = doc.add_paragraph()
        
        run = para.add_run()
        run.add_picture(image_stream, width=Inches(width_inches))
        
        result = upload_docx_to_blob(doc, blob_uri)
        return json.dumps({
            "success": True,
            "message": "Image inserted successfully",
            "blob_uri": blob_uri
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)

@server.tool()
async def format_document(
    blob_uri: str,
    paragraph_index: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    font_color_rgb: Optional[List[int]] = None,
    alignment: Optional[str] = None,
    left_margin_inches: Optional[float] = None,
    right_margin_inches: Optional[float] = None,
    top_margin_inches: Optional[float] = None,
    bottom_margin_inches: Optional[float] = None
) -> str:
    """Apply formatting to text and page margins in the document.

    Args:
        blob_uri: Azure Blob Storage URI of the document
        paragraph_index: Index of paragraph to format (None for all paragraphs)
        bold: Make text bold
        italic: Make text italic
        underline: Underline text
        font_name: Font name (e.g., "Arial", "Times New Roman")
        font_size: Font size in points
        font_color_rgb: Font color as RGB array [R, G, B] (0-255)
        alignment: Text alignment ("left", "center", "right", "justify")
        left_margin_inches: Left page margin in inches
        right_margin_inches: Right page margin in inches
        top_margin_inches: Top page margin in inches
        bottom_margin_inches: Bottom page margin in inches
    
    Returns:
        Success message
    """
    try:
        doc = download_docx_from_blob(blob_uri)
        
        # Apply paragraph formatting
        if any([bold is not None, italic is not None, underline is not None, 
                font_name, font_size, font_color_rgb, alignment]):
            
            paragraphs_to_format = []
            if paragraph_index is not None:
                if paragraph_index < len(doc.paragraphs):
                    paragraphs_to_format = [doc.paragraphs[paragraph_index]]
            else:
                paragraphs_to_format = doc.paragraphs
            
            for para in paragraphs_to_format:
                # Apply alignment
                if alignment:
                    alignment_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if alignment.lower() in alignment_map:
                        para.alignment = alignment_map[alignment.lower()]
                
                # Apply run formatting
                for run in para.runs:
                    if bold is not None:
                        run.bold = bold
                    if italic is not None:
                        run.italic = italic
                    if underline is not None:
                        run.underline = underline
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = Pt(font_size)
                    if font_color_rgb and len(font_color_rgb) == 3:
                        run.font.color.rgb = RGBColor(
                            font_color_rgb[0], 
                            font_color_rgb[1], 
                            font_color_rgb[2]
                        )
        
        # Apply page margin formatting
        if any([left_margin_inches, right_margin_inches, top_margin_inches, bottom_margin_inches]):
            for section in doc.sections:
                if left_margin_inches is not None:
                    section.left_margin = Inches(left_margin_inches)
                if right_margin_inches is not None:
                    section.right_margin = Inches(right_margin_inches)
                if top_margin_inches is not None:
                    section.top_margin = Inches(top_margin_inches)
                if bottom_margin_inches is not None:
                    section.bottom_margin = Inches(bottom_margin_inches)
        
        result = upload_docx_to_blob(doc, blob_uri)
        return json.dumps({
            "success": True,
            "message": "Formatting applied successfully",
            "blob_uri": blob_uri
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, ensure_ascii=False)
